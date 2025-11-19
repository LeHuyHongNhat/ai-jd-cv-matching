"""
Script để tạo sample CV và JD để test
Sử dụng dữ liệu mẫu từ tests/test_data.py
"""
import os
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


def create_sample_cv_pdf(filename: str, content: str):
    """Tạo file PDF mẫu cho CV"""
    print(f"Creating sample CV: {filename}")
    
    # Tạo PDF
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Chia content thành các đoạn
    paragraphs = content.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            p = Paragraph(para.replace('\n', '<br/>'), styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 12))
    
    doc.build(story)
    print(f"✓ Created: {filename}")


def create_sample_jd_docx(filename: str, content: str):
    """Tạo file DOCX mẫu cho JD"""
    print(f"Creating sample JD: {filename}")
    
    doc = Document()
    doc.add_heading('Job Description', 0)
    
    # Chia content thành các đoạn
    paragraphs = content.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    
    doc.save(filename)
    print(f"✓ Created: {filename}")


def main():
    """Tạo sample files"""
    from tests.test_data import SAMPLE_CV_TEXT, SAMPLE_JD_TEXT
    
    # Tạo thư mục samples
    samples_dir = Path("samples")
    samples_dir.mkdir(exist_ok=True)
    
    print("=" * 60)
    print("Creating sample files for demo")
    print("=" * 60)
    print()
    
    # Tạo sample CV (PDF)
    cv_path = samples_dir / "sample_cv_1.pdf"
    create_sample_cv_pdf(str(cv_path), SAMPLE_CV_TEXT)
    
    # Tạo thêm 2 CV nữa với nội dung tương tự
    cv2_text = SAMPLE_CV_TEXT.replace("Le Huy Hong Nhat", "Nguyen Van A")
    cv_path2 = samples_dir / "sample_cv_2.pdf"
    create_sample_cv_pdf(str(cv_path2), cv2_text)
    
    cv3_text = SAMPLE_CV_TEXT.replace("Le Huy Hong Nhat", "Tran Thi B")
    cv_path3 = samples_dir / "sample_cv_3.pdf"
    create_sample_cv_pdf(str(cv_path3), cv3_text)
    
    # Tạo sample JD (DOCX)
    jd_path = samples_dir / "sample_jd.docx"
    create_sample_jd_docx(str(jd_path), SAMPLE_JD_TEXT)
    
    print()
    print("=" * 60)
    print("Sample files created successfully!")
    print("=" * 60)
    print()
    print("Files created:")
    print(f"  • {cv_path}")
    print(f"  • {cv_path2}")
    print(f"  • {cv_path3}")
    print(f"  • {jd_path}")
    print()
    print("To run demo:")
    print(f"  python demo_simple.py {cv_path} {jd_path}")
    print(f"  python demo_matching.py {samples_dir} {jd_path}")


if __name__ == "__main__":
    main()

